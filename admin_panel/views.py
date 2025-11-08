from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.hashers import check_password
from django.http import FileResponse, JsonResponse
from django.db.models import Sum
from django.db.models.functions import ExtractMonth, ExtractYear
from .models import Administrador
from registros_crud.models import *
from clientes_crud.models import *
from operadores_crud.models import *
from .utils import admin_required

from django.utils import timezone
from datetime import datetime, timedelta
import locale
from django.conf import settings
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os, time, threading
from docx2pdf import convert


def index(request):
    if request.session.get("admin_id"):
        return redirect("admin_panel")
    if request.session.get("operador_id"):
        return redirect("operador_panel")

    if request.method == "POST":
        rut = request.POST.get("rut", "").strip()
        pwd = request.POST.get("password", "")

        try:
            admin = Administrador.objects.get(rut=rut)
            if check_password(pwd, admin.contraseña):
                request.session["admin_id"] = admin.id
                request.session["admin_nombre"] = admin.nombre
                return redirect("admin_panel")
        except Administrador.DoesNotExist:
            pass

        try:
            operador = Operador.objects.get(rut=rut)
            if check_password(pwd, operador.contraseña):
                if not operador.estado:
                    messages.error(request, "Su cuenta ha sido bloqueada, hable con un administrador para resolver la situación.")
                    return redirect("inicio_sesion")
                request.session["operador_id"] = operador.id
                request.session["operador_nombre"] = operador.nombre
                return redirect("operador_panel")
        except Operador.DoesNotExist:
            pass

    return render(request, "admin_panel/inicio_sesion.html")

@admin_required
def admin_panel(request):
    registros_info = Residuos.objects.all().select_related('idOperador','idCliente','idUbicacion').order_by('-id')
    clientes_info = Cliente.objects.order_by('nombre')
    operadores_info = Operador.objects.order_by('nombre')
    ubicaciones_info = UbicacionCl.objects.order_by('calle')

    ahora = datetime.now()
    mes_actual = ahora.month
    anio_actual = ahora.year

    registros_mes = Residuos.objects.filter(
        fechaRegistro__year=anio_actual,
        fechaRegistro__month=mes_actual
    )

    campos_residuos = ['plastico', 'papel', 'carton', 'film', 'latas', 'palets', 'chatarra', 'vidrio', 'tetrapack']
    total_mes_kg = 0
    desecho_mas_cantidad = None
    max_valor = 0

    for campo in campos_residuos:
        suma = registros_mes.aggregate(total=Sum(campo))['total'] or 0
        total_mes_kg += suma
        if suma > max_valor:
            max_valor = suma
            desecho_mas_cantidad = campo.capitalize()

    cantidad_registros_mes = registros_mes.count()

    return render(request, "admin_panel/admin_panel.html", {
        "active": "panel",
        'registros_info': registros_info,
        'clientes_info': clientes_info,
        'operadores_info': operadores_info,
        'ubicaciones_info': ubicaciones_info,
        'total_mes_kg': total_mes_kg,
        'cantidad_registros_mes': cantidad_registros_mes,
        'desecho_mas_cantidad': desecho_mas_cantidad,
    })

def logout_admin(request):
    request.session.flush()
    return redirect("inicio_sesion")

@admin_required
def generar_certificado(request):
    clientes_info = Cliente.objects.order_by('nombre')
    return render(request, "admin_panel/generar_certificado.html",{
        'clientes_info': clientes_info,
    })

def procesarYResponderCertificado(doc, nombre_docx, nombre_pdf, request):
    import os
    import time
    import pythoncom
    from django.http import FileResponse
    from django.utils import timezone

    temp_dir = os.path.join(settings.BASE_DIR, 'temp_certs')
    os.makedirs(temp_dir, exist_ok=True)

    ruta_docx = os.path.join(temp_dir, nombre_docx)
    ruta_pdf = os.path.join(temp_dir, nombre_pdf)
    doc.save(ruta_docx)

    try:
        pythoncom.CoInitialize()
        convert(ruta_docx, ruta_pdf)
        pythoncom.CoUninitialize()
    except Exception as e:
        messages.error(request, f"Error al convertir a PDF: {e}")
        return redirect('generar_certificado')

    def borrar_temporales():
        time.sleep(10)
        for f in [ruta_docx, ruta_pdf]:
            try:
                os.remove(f)
            except Exception:
                pass

    try:
        pdf = open(ruta_pdf, 'rb')
        response = FileResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{nombre_pdf}"'
        threading.Thread(target=borrar_temporales).start()
        return response
    except Exception as e:
        messages.error(request, f"Error inesperado: {e}")
        return redirect('generar_certificado')

def generarCertificado(request):
    if not request.session.get('admin_id'):
        messages.warning(request, "Debes iniciar sesión para generar certificados.")
        return redirect('admin_panel')

    if request.method != 'POST':
        messages.error(request, "Método no permitido.")
        return redirect('generar_certificado')

    cliente_id = request.POST.get('cliente_id')
    anio = request.POST.get('anio')
    mes = request.POST.get('mes')
    tipo = request.POST.get('tipo_certificado')

    if not all([cliente_id, anio, mes, tipo]):
        messages.error(request, "Completa todos los campos para generar el certificado.")
        return redirect('generar_certificado')

    try:
        cliente_id = int(cliente_id)
        anio = int(anio)
        mes = int(mes)
    except ValueError:
        messages.error(request, "Datos de cliente, año o mes inválidos.")
        return redirect('generar_certificado')

    cliente = Cliente.objects.get(id=cliente_id)
    residuos = Residuos.objects.filter(idCliente_id=cliente_id, fechaRegistro__year=anio, fechaRegistro__month=mes)

    try:
        locale.setlocale(locale.LC_TIME, 'es_CL.utf8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Chile.1252')
        except locale.Error:
            locale.setlocale(locale.LC_TIME, '')

    fecha_actual = datetime.now().strftime("%A, %d DE %B %Y").upper()
    nombre_mes = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"][mes - 1]

    if tipo == "2":
        eco_tipos = {
            'pallet1': 'palets',
            'carton1': 'carton',
            'papel1': 'papel',
            'plastico1': 'plastico',
            'aluminio1': 'latas',
            'tetrapak1': 'tetrapack'
        }
        totales = {clave: 0.0 for clave in eco_tipos.keys()}
        total_vidrio = 0.0

        for r in residuos:
            for clave, campo in eco_tipos.items():
                valor = getattr(r, campo)
                if valor:
                    totales[clave] += float(valor)
            if hasattr(r, 'vidrio') and r.vidrio is not None:
                total_vidrio += float(r.vidrio)

        total_plastico = totales.get('plastico1', 0.0)
        total_carton = totales.get('carton1', 0.0)
        total_papel = totales.get('papel1', 0.0)
        total_palets = totales.get('pallet1', 0.0)

        com1 = round((total_plastico * 1000 / 500) + (total_vidrio * 1000 / 500))
        ene1 = round((total_plastico * 5) + ((total_carton + total_papel) * 7000 / 1000) + (total_vidrio * 670 / 1000))
        emi1 = round((total_plastico * 2.5) + ((total_carton + total_papel) * 7000 / 1000) + (total_vidrio * 1300 / 100))
        agua1 = round((total_plastico * 36.26) + ((total_carton + total_papel) * 270000 / 1000))
        arb_calc = (total_palets * 1 / 10) + ((total_carton + total_papel) * 17 / 1000)
        arb1 = round(arb_calc) if arb_calc >= 1 else 0

        total_eco = sum(totales.values()) + total_vidrio

        ruta_plantilla = os.path.join(settings.MEDIA_ROOT, 'plantillas', 'plantillaEco.docx')
        doc = Document(ruta_plantilla)

        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for clave, valor in totales.items():
                        if clave in celda.text:
                            celda.text = celda.text.replace(clave, f"{valor:.1f}")
                    if "vidrio1" in celda.text:
                        celda.text = celda.text.replace("vidrio1", f"{total_vidrio:.1f}")
                    if "Total1" in celda.text:
                        celda.text = celda.text.replace("Total1", f"{total_eco:.1f}")

        for p in doc.paragraphs:
            for run in p.runs:
                if "FechaCreacion" in run.text:
                    run.text = run.text.replace("FechaCreacion", fecha_actual)
                if "Cliente1" in run.text:
                    run.text = run.text.replace("Cliente1", cliente.nombre)
                if "Mes1" in run.text:
                    run.text = run.text.replace("Mes1", nombre_mes)
                for clave, valor in {"com1": com1, "ene1": ene1, "emi1": emi1, "agua1": agua1, "arb1": arb1}.items():
                    if clave in run.text:
                        run.text = run.text.replace(clave, str(valor))
                        run.font.name = 'Calibri'
                        run.bold = True
                        run.font.size = Pt(19)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        return procesarYResponderCertificado(doc, f"TEMP_ECO_{cliente_id}_{anio}_{mes}_{timezone.now().timestamp()}.docx", f"Certificado_Eco_{cliente_id}_{anio}_{mes}.pdf", request)

    # Certificado tipo 1 (CTTR)
    tipos_residuos = ["plastico", "papel", "carton", "film", "latas", "palets", "chatarra", "vidrio", "tetrapack"]
    totales = {tipo: 0.0 for tipo in tipos_residuos}

    for r in residuos:
        for tipo in tipos_residuos:
            valor = getattr(r, tipo)
            if valor:
                totales[tipo] += float(valor)

    resumen = [f"{tipo.capitalize()}: {totales[tipo]:.1f}" for tipo in tipos_residuos if totales[tipo] > 0]
    if not resumen:
        messages.warning(request, "No se encontraron registros para el cliente y mes seleccionado.")
        return redirect('generar_certificado')

    texto = "\n".join(resumen)
    total_kg = sum(totales.values())

    ruta_plantilla = os.path.join(settings.MEDIA_ROOT, 'plantillas', 'plantillaCttr.docx')
    doc = Document(ruta_plantilla)
    fuente = "Carmen Sans SemiBold"

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if "Residuos1" in celda.text:
                    celda.text = ""
                    p = celda.paragraphs[0]
                    run = p.add_run(texto)
                    run.font.name = fuente
                    run.font.size = Pt(8)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), fuente)
                elif "Total1" in celda.text:
                    celda.text = ""
                    p = celda.paragraphs[0]
                    run = p.add_run(f"{total_kg:.1f} kg")
                    run.font.name = fuente
                    run.font.size = Pt(8)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), fuente)
                elif "Cliente1" in celda.text:
                    celda.text = ""
                    p = celda.paragraphs[0]
                    run = p.add_run(cliente.nombre)
                    run.font.name = fuente
                    run.font.size = Pt(8)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), fuente)

    for p in doc.paragraphs:
        if "FechaCreacion" in p.text:
            p.text = p.text.replace("FechaCreacion", fecha_actual)
        if "Mes1" in p.text:
            p.text = p.text.replace("Mes1", nombre_mes)
        if "Anio1" in p.text:
            p.text = p.text.replace("Anio1", str(anio))

    return procesarYResponderCertificado(doc, f"TEMP_CTTR_{cliente_id}_{anio}_{mes}_{timezone.now().timestamp()}.docx", f"Certificado_CTTR_{cliente_id}_{anio}_{mes}.pdf", request)

def api_grafico_barras(request):
    """API que devuelve datos agregados para el gráfico de barras"""
    
    
    anio = request.GET.get('anio')
    mes = request.GET.get('mes')
    cliente_id = request.GET.get('cliente_id')
    residuos_seleccionados = request.GET.get('residuos', '').split(',')
    
    if not anio:
        anio = datetime.now().year
    else:
        anio = int(anio)
    
    query = Residuos.objects.filter(fechaRegistro__year=anio)
    
    if mes:
        query = query.filter(fechaRegistro__month=int(mes))
    
    if cliente_id:
        query = query.filter(idCliente_id=int(cliente_id))
    
    mapeo_nombres = {
        'plastico': 'Plástico',
        'papel': 'Papel',
        'carton': 'Cartón',
        'film': 'Films',
        'latas': 'Latas',
        'palets': 'Palets',
        'chatarra': 'Chatarra',
        'vidrio': 'Vidrio',
        'tetrapack': 'Tetrapack',
    }
    
    labels = []
    values = []
    
    for residuo in residuos_seleccionados:
        if residuo in mapeo_nombres:
            total = query.aggregate(suma=Sum(residuo))['suma']
            labels.append(mapeo_nombres[residuo])
            values.append(float(total or 0))
    
    meses_nombres = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", 
                     "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    
    titulo_partes = []
    if cliente_id:
        try:
            cliente = Cliente.objects.get(id=cliente_id)
            titulo_partes.append(f"Cliente: {cliente.nombre}")
        except Cliente.DoesNotExist:
            pass
    
    if mes:
        titulo_partes.append(f"{meses_nombres[int(mes)]} {anio}")
    else:
        titulo_partes.append(f"Año {anio}")
    
    titulo = " - ".join(titulo_partes) if titulo_partes else f"Residuos {anio}"
    
    return JsonResponse({
        'labels': labels,
        'values': values,
        'titulo': titulo
    })